import sys
import json
import os
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def generate_exact_bast_docx(template_docx_path, output_docx_path, data_json_path):
    with open(data_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    nomor_bast = data.get('nomor_bast', '')
    nomor_spk = data.get('nomor_spk', '')
    nama_mitra = data.get('nama_mitra', '')
    pekerjaan_mitra = data.get('pekerjaan_mitra', 'Mitra BPS')
    alamat_mitra = data.get('alamat_mitra', '')
    tahun = str(data.get('tahun', '2025'))
    
    # Tanggal pembuka BAST
    hari = data.get('hari', 'Jumat')
    tanggal_teks = data.get('tanggal_teks', 'tiga puluh satu')
    bulan_teks = data.get('bulan_teks', 'Maret')
    tahun_teks = data.get('tahun_teks', 'dua ribu dua puluh lima')
    tanggal_angka = data.get('tanggal_angka', '31 Maret 2025')
    tanggal_spk_text = data.get('tanggal_spk_text', '')

    doc = docx.Document(template_docx_path)

    # 1. In-place Run Injection
    for p in doc.paragraphs:
        txt = p.text

        # Update Judul Tahun (P1)
        if 'KEGIATAN SURVEI/SENSUS TAHUN' in txt:
            for r in p.runs:
                if '2023' in r.text or '2024' in r.text or '2025' in r.text:
                    r.text = r.text.replace('2023', tahun).replace('2024', tahun).replace('2025', tahun)

        # Update Nomor BAST (P2)
        if 'Nomor:' in txt or 'NOMOR:' in txt:
            # Clear trailing empty runs and insert nomor into run 1
            if len(p.runs) > 1:
                p.runs[1].text = nomor_bast
                for r in p.runs[2:]:
                    r.text = ''
            else:
                p.text = f"Nomor: {nomor_bast}"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Bookman Old Style'
                    r.font.size = Pt(11)

        # Update Tanggal Pembuka (P4)
        if 'Pada hari ini' in txt:
            # In-place run replacement
            # P4 runs: ['Pada hari ini ', 'Jumat', ', ', 'tanggal ', 'tiga puluh satu', ' bulan ', 'Maret', ' ta', 'hun ', 'dua ribu dua puluh tiga', ' ', '(', '31 Maret', ' 2023', '), bertempat di Kantor BPS Kabupaten ', 'Tasikmalaya', ' dengan alamat Jalan R.A.A Kusumahsubrata Komplek Perkantoran Kertasari, Kabupaten ', 'Tasikmalaya', ', yang bertanda tangan di bawah ini:']
            for r in p.runs:
                if 'Jum' in r.text:
                    r.text = hari
                elif r.text == 'tiga puluh satu':
                    r.text = tanggal_teks
                elif r.text == 'Maret':
                    r.text = bulan_teks
                elif r.text == 'dua ribu dua puluh tiga':
                    r.text = tahun_teks
                elif r.text == '31 Maret':
                    r.text = tanggal_angka.rsplit(' ', 1)[0]
                elif r.text == ' 2023':
                    r.text = ' ' + tahun

        # Update PIHAK KEDUA (P7)
        if txt.startswith('2.') or 'PIHAK KEDUA' in txt and ('berkedudukan di' in txt or 'bertindak untuk' in txt):
            # Clear all existing runs/fields from XML paragraph
            pPr = p._p.pPr
            for child in list(p._p):
                if child.tag.endswith('r'):
                    p._p.remove(child)

            # Re-create clean runs without any MERGEFIELD artifacts
            r0 = p.add_run('2. ')
            r0.font.name = 'Bookman Old Style'
            r0.font.size = Pt(11)

            r1 = p.add_run(nama_mitra)
            r1.font.name = 'Bookman Old Style'
            r1.font.size = Pt(11)

            r2 = p.add_run('\t:\t')
            r2.font.name = 'Bookman Old Style'
            r2.font.size = Pt(11)

            r3 = p.add_run(f"{pekerjaan_mitra}, berkedudukan di {alamat_mitra}, bertindak untuk dan atas nama diri sendiri, selanjutnya disebut ")
            r3.font.name = 'Bookman Old Style'
            r3.font.size = Pt(11)

            r4 = p.add_run('PIHAK KEDUA')
            r4.font.name = 'Bookman Old Style'
            r4.font.size = Pt(11)
            r4.font.bold = True

            r5 = p.add_run('.')
            r5.font.name = 'Bookman Old Style'
            r5.font.size = Pt(11)

        # Update Rujukan Perjanjian Kerja SPK (P9)
        if 'Berdasarkan' in txt and 'Perjanjian Kerja' in txt:
            # Clear all existing runs/fields from XML paragraph
            for child in list(p._p):
                if child.tag.endswith('r'):
                    p._p.remove(child)

            tgl_ref = f", tanggal {tanggal_spk_text}" if tanggal_spk_text else ""
            r_spk = p.add_run(f"Berdasarkan Perjanjian Kerja Nomor {nomor_spk}{tgl_ref}, bersama ini PIHAK KEDUA telah menyerahkan pekerjaan kepada PIHAK PERTAMA, dengan ketentuan sebagai berikut:")
            r_spk.font.name = 'Bookman Old Style'
            r_spk.font.size = Pt(11)

    # 2. Update Tanda Tangan (Table 0)
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # Kolom 0: PIHAK KEDUA, Kolom 1: PIHAK PERTAMA
        if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
            c0 = table.rows[0].cells[0]
            c1 = table.rows[0].cells[1]

            c0.text = f"PIHAK KEDUA\n\n\n\n\n{nama_mitra}"
            c1.text = f"PIHAK PERTAMA\n\n\n\n\nDindin Muldiana, S.ST. MP."

            for p in c0.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Bookman Old Style'
                    r.font.size = Pt(11)
                    if nama_mitra in r.text:
                        r.font.bold = True
                        r.font.underline = True

            for p in c1.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Bookman Old Style'
                    r.font.size = Pt(11)
                    if 'Dindin Muldiana' in r.text:
                        r.font.bold = True
                        r.font.underline = True

    doc.save(output_docx_path)

    # Convert to PDF if --pdf flag passed
    if len(sys.argv) >= 5 and sys.argv[4] == '--pdf':
        pdf_out_path = output_docx_path.rsplit('.', 1)[0] + '.pdf'
        try:
            import win32com.client, pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            abs_docx = os.path.abspath(output_docx_path)
            abs_pdf = os.path.abspath(pdf_out_path)
            doc_app = word.Documents.Open(abs_docx)
            doc_app.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
            doc_app.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            print(f"SUCCESS_PDF:{abs_pdf}")
        except Exception as e:
            print(f"PDF_CONVERT_ERROR:{str(e)}")

if __name__ == '__main__':
    if len(sys.argv) >= 4:
        generate_exact_bast_docx(sys.argv[1], sys.argv[2], sys.argv[3])

import os
import sys
import json
import docx
import win32com.client
import pythoncom

def generate_exact_spk_pdf(input_docx_path, output_pdf_path, data_json_path):
    with open(data_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = docx.Document(input_docx_path)

    nomor_spk = data.get('nomor_spk', '')
    nama_mitra = data.get('nama_mitra', '')
    pekerjaan_mitra = data.get('pekerjaan_mitra', '')
    alamat_mitra = data.get('alamat_mitra', '')
    periode_label = data.get('periode_label', '')
    total_honor = data.get('total_honor', '')
    terbilang_honor = data.get('terbilang_honor', '')
    tahun = data.get('tahun', '')
    items = data.get('items', [])

    # 1. PARAGRAPH REPLACEMENTS
    for p in doc.paragraphs:
        if '1001/PPK/SPK/03/2024' in p.text:
            for r in p.runs:
                if '1001/PPK/SPK/03/2024' in r.text:
                    r.text = r.text.replace('1001/PPK/SPK/03/2024', nomor_spk)
        if 'TAHUN 2024' in p.text:
            for r in p.runs:
                if 'TAHUN 2024' in r.text:
                    r.text = r.text.replace('TAHUN 2024', f'TAHUN {tahun}')
        if 'LINA KARLINA' in p.text:
            for r in p.runs:
                if 'LINA KARLINA' in r.text:
                    r.text = r.text.replace('LINA KARLINA', nama_mitra)
        if 'Lainnya/ Belum Bekerja' in p.text:
            for r in p.runs:
                if 'Lainnya/ Belum Bekerja' in r.text:
                    r.text = r.text.replace('Lainnya/ Belum Bekerja', pekerjaan_mitra)
        if 'Kp. Pameungpeuk RT/RW : 24/03 Desa Sukarasa Kec. Salawu' in p.text:
            for r in p.runs:
                if 'Kp. Pameungpeuk RT/RW : 24/03 Desa Sukarasa Kec. Salawu' in r.text:
                    r.text = r.text.replace('Kp. Pameungpeuk RT/RW : 24/03 Desa Sukarasa Kec. Salawu', alamat_mitra)
        if '1 Maret 2024 sampai dengan tanggal 31 Maret 2024' in p.text:
            for r in p.runs:
                if '1 Maret 2024 sampai dengan tanggal 31 Maret 2024' in r.text:
                    r.text = r.text.replace('1 Maret 2024 sampai dengan tanggal 31 Maret 2024', periode_label)
        if 'Rp. 1.160.000 (Satu Juta Seratus Enam Puluh Ribu Rupiah)' in p.text:
            for r in p.runs:
                if 'Rp. 1.160.000 (Satu Juta Seratus Enam Puluh Ribu Rupiah)' in r.text:
                    r.text = r.text.replace('Rp. 1.160.000 (Satu Juta Seratus Enam Puluh Ribu Rupiah)', f'{total_honor} ({terbilang_honor})')
        if 'Rp. 1.160.000,00 (Satu Juta Seratus Enam Puluh Ribu Rupiah)' in p.text:
            for r in p.runs:
                if 'Rp. 1.160.000,00 (Satu Juta Seratus Enam Puluh Ribu Rupiah)' in r.text:
                    r.text = r.text.replace('Rp. 1.160.000,00 (Satu Juta Seratus Enam Puluh Ribu Rupiah)', f'{total_honor},00 ({terbilang_honor})')

    # 2. SIGNATURE TABLE (TABLE 0)
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        for row in t0.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if 'LINA KARLINA' in p.text:
                        for r in p.runs:
                            if 'LINA KARLINA' in r.text:
                                r.text = r.text.replace('LINA KARLINA', nama_mitra)

    # 3. ANNEX TABLE (TABLE 1)
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        
        # Populate activity items
        for i in range(8):
            row_idx = 3 + i
            if row_idx < len(t1.rows):
                row = t1.rows[row_idx]
                if i < len(items):
                    item = items[i]
                    # col 0: No
                    # col 1: Uraian Tugas
                    # col 2: Jangka Waktu
                    # col 3: Volume
                    # col 4: Satuan
                    # col 5: Harga Satuan
                    # col 6: Nilai Perjanjian
                    # col 7: Beban Anggaran
                    if len(row.cells) >= 8:
                        row.cells[1].text = item.get('nama', '')
                        row.cells[2].text = item.get('periode', '')
                        row.cells[3].text = str(item.get('volume', 1))
                        row.cells[4].text = item.get('satuan', 'dokumen')
                        row.cells[5].text = item.get('harga_satuan', '')
                        row.cells[6].text = item.get('nilai_perjanjian', '')
                        row.cells[7].text = item.get('mak', '')
                else:
                    if len(row.cells) >= 8:
                        row.cells[1].text = ''
                        row.cells[2].text = ''
                        row.cells[3].text = ''
                        row.cells[4].text = ''
                        row.cells[5].text = 'Rp. 00'
                        row.cells[6].text = 'Rp. 00'
                        row.cells[7].text = ''

        # Total Row (Row 11)
        if len(t1.rows) > 11:
            total_row = t1.rows[11]
            if len(total_row.cells) >= 7:
                total_row.cells[0].text = f'Terbilang: {terbilang_honor}'
                total_row.cells[6].text = f'{total_honor}, 00'

    # Save to temp docx
    temp_docx = output_pdf_path.replace('.pdf', '_temp.docx')
    doc.save(temp_docx)

    # Convert to PDF via MS Word COM API
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        wdoc = word.Documents.Open(os.path.abspath(temp_docx))
        wdoc.SaveAs(os.path.abspath(output_pdf_path), FileFormat=17) # 17 = wdFormatPDF
        wdoc.Close()
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

    if os.path.exists(temp_docx):
        try:
            os.remove(temp_docx)
        except:
            pass

if __name__ == '__main__':
    if len(sys.argv) >= 4:
        generate_exact_spk_pdf(sys.argv[1], sys.argv[2], sys.argv[3])

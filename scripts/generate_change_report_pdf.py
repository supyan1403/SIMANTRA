import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client
import pythoncom

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_report_docx(output_docx_path):
    doc = docx.Document()

    # Set Margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main_title = p_title.add_run("LAPORAN DOKUMENTASI PEMBARUAN SISTEM\nSIMANTRA BPS KABUPATEN TASIKMALAYA")
    r_main_title.font.name = "Arial"
    r_main_title.font.size = Pt(15)
    r_main_title.font.bold = True
    r_main_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Rekapitulasi Pengembangan Fitur, Logika SBML, Otomasi SPK, dan UI\nPeriode: 28 Agustus 2026 – 29 Agustus 2026")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(10)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    # Horizontal Divider Line Table
    div_table = doc.add_table(rows=1, cols=1)
    div_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    div_cell = div_table.cell(0, 0)
    set_cell_background(div_cell, "2563EB")
    div_cell.height = Inches(0.04)
    p_div = div_cell.paragraphs[0]
    p_div.paragraph_format.space_before = Pt(0)
    p_div.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. EXECUTIVE SUMMARY
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Ringkasan Eksekutif (Executive Summary)")
    r_h1.font.name = "Arial"
    r_h1.font.color.rgb = RGBColor(30, 58, 138)

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.line_spacing = 1.15
    p_exec.paragraph_format.space_after = Pt(6)
    p_exec.add_run(
        "Dalam kurun waktu pengembangan intensif (28–29 Agustus 2026), aplikasi SIMANTRA (Sistem Informasi Monitoring Alokasi Pekerjaan dan Honor Mitra BPS) "
        "telah mengalami pembaruan menyeluruh pada 4 pilar arsitektur utama: otomasi penomoran SPK/BAST, implementasi aturan batas Standar Biaya Masukan Lainnya (SBML) multi-tugas yang akurat sesuai data riil BPS, "
        "desain antarmuka sidebar anti-reflow (zero-shift), serta sinkronisasi penuh basis data 2.566 mitra statistik ke GitHub."
    )

    # 2. DETAIL PEMBARUAN PER MODUL
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Rincian Pembaruan Fitur & Logika Sistem")
    r_h2.font.name = "Arial"
    r_h2.font.color.rgb = RGBColor(30, 58, 138)

    # Table of Features
    table_feat = doc.add_table(rows=1, cols=3)
    table_feat.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_feat.rows[0].cells
    hdr_cells[0].text = "Modul / Fitur"
    hdr_cells[1].text = "Sebelum Pembaruan"
    hdr_cells[2].text = "Hasil Pembaruan Terbaru (Hari Ini)"

    for c in hdr_cells:
        set_cell_background(c, "1E293B")
        set_cell_margins(c, 100, 100, 140, 140)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

    features_data = [
        (
            "Penomoran SPK & BAST (/spk/penomoran)",
            "Pola penomoran statis, counter nomor bercampur antar-format, dropdown kegiatan terbatas dan tidak memiliki fitur pencarian.",
            "• Searchable Dropdown dinamis dengan 49 kegiatan statistik BPS.\n• Pengelolaan pola fleksibel (Tambah, Edit via AJAX, Reset Standar BPS).\n• Counter nomor otomatis independen per kegiatan & per tahun anggaran.\n• Live Preview nomor SPK/BAST pertama secara instan."
        ),
        (
            "Logika & Validasi Limit SBML",
            "Batas SBML tunggal (flat Rp 4.500.000,- / Rp 7.500.000,-) tanpa membedakan tugas pendataan lapangan dan pengolahan.",
            "• Pemisahan Standar SBML Riil BPS:\n  - Th 2024: Lapangan Rp 3.326.000,- | Data Rp 3.077.000,- | Gabungan Rp 6.403.000,-\n  - Th 2025: Lapangan Rp 4.500.000,- | Data Rp 3.000.000,- | Gabungan Rp 7.500.000,-\n• Evaluasi cerdas status honor bulanan kumulatif (Sesuai vs Melebihi SBML)."
        ),
        (
            "Master SBML & Badge Kegiatan",
            "Tabel Master SBML kosong dan tidak ada indikator jenis tugas pada nama kegiatan di tabel alokasi.",
            "• Modul Master SBML (/master-sbml) lengkap dengan rincian kolom Pencacahan, Pengolahan, dan kalkulator total gabungan.\n• Badge visual otomatis pada Monitoring & Master Kegiatan:\n  - [📍 Pencacahan] (Biru) untuk survei/lapangan.\n  - [💻 Pengolahan] (Kuning Amber) untuk pengolahan/entri data."
        ),
        (
            "Antarmuka Sidebar (UI/UX)",
            "Saat sidebar dibuka/ditutup, ikon bergeser (shift 2-4px), logo terpotong saat collapse, dan footer profil melompat ke bawah.",
            "• Sumbu Jangkar Tetap (36px Fixed Axis Lock) - Ikon 100% diam tanpa bergeser.\n• Clipping Box Architecture (bebas reflow/jumping).\n• Header brand dilengkapi margin atas lapang & logo kendi utuh simetris.\n• Footer profil 1-baris permanen."
        ),
        (
            "Basis Data & Kolaborasi Tim",
            "Database kosong setelah clone dan membutuhkan import Excel manual yang rawan gagal.",
            "• database/database.sqlite lengkap (2.566 Mitra, 47 Kegiatan, 737 Alokasi, SBML 2024/2025) terintegrasi langsung di repositori.\n• Rekan kerja cukup menjalankan 'git pull' & 'php artisan serve' tanpa konfigurasi manual."
        )
    ]

    for mod, bef, aft in features_data:
        row_cells = table_feat.add_row().cells
        row_cells[0].text = mod
        row_cells[1].text = bef
        row_cells[2].text = aft
        for i, c in enumerate(row_cells):
            set_cell_background(c, "F8FAFC" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(c, 90, 90, 120, 120)
            p = c.paragraphs[0]
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. STANDAR ANGKA SBML RESMI BPS KAB. TASIKMALAYA
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Tabel Standar Batas SBML Resmi BPS")
    r_h3.font.name = "Arial"
    r_h3.font.color.rgb = RGBColor(30, 58, 138)

    p_sbml_desc = doc.add_paragraph()
    p_sbml_desc.add_run("Berikut adalah batas honor bulanan per mitra berdasarkan Standar Biaya Masukan Lainnya (SBML) yang diintegrasikan ke dalam sistem:")

    table_sbml = doc.add_table(rows=1, cols=4)
    table_sbml.alignment = WD_TABLE_ALIGNMENT.CENTER
    sbml_hdr = table_sbml.rows[0].cells
    sbml_hdr[0].text = "Tahun Anggaran"
    sbml_hdr[1].text = "Batas Pencacahan (Lapangan)"
    sbml_hdr[2].text = "Batas Pengolahan (Data)"
    sbml_hdr[3].text = "Total Batas Maksimal Gabungan"

    for c in sbml_hdr:
        set_cell_background(c, "0284C7")
        set_cell_margins(c, 100, 100, 140, 140)
        p = c.paragraphs[0]
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

    sbml_rows = [
        ("Tahun 2024 (Data Riil Excel BPS)", "Rp 3.326.000,-", "Rp 3.077.000,-", "Rp 6.403.000,-"),
        ("Tahun 2025 (Standar SBM Nasional)", "Rp 4.500.000,-", "Rp 3.000.000,-", "Rp 7.500.000,-")
    ]

    for yr, pnc, png, tot in sbml_rows:
        rc = table_sbml.add_row().cells
        rc[0].text = yr
        rc[1].text = pnc
        rc[2].text = png
        rc[3].text = tot
        for c in rc:
            set_cell_background(c, "F0F9FF")
            set_cell_margins(c, 80, 80, 120, 120)
            p = c.paragraphs[0]
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. INSTRUKSI SETUP UNTUK REKAN KERJA
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Panduan Eksekusi Setup Rekan Kerja (Instalasi 1 Langkah)")
    r_h4.font.name = "Arial"
    r_h4.font.color.rgb = RGBColor(30, 58, 138)

    p_setup = doc.add_paragraph()
    p_setup.add_run(
        "Bagi rekan kerja yang melakukan clone repositori baru di komputernya, cukup menjalankan urutan perintah berikut di terminal/PowerShell:\n"
    )

    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_cell = box_table.cell(0, 0)
    set_cell_background(b_cell, "0F172A")
    set_cell_margins(b_cell, 120, 120, 160, 160)
    p_code = b_cell.paragraphs[0]
    p_code.paragraph_format.line_spacing = 1.15
    r_code = p_code.add_run(
        "git clone https://github.com/supyan1403/SIMANTRA.git\n"
        "cd SIMANTRA\n"
        "composer install --ignore-platform-reqs\n"
        "copy .env.example .env\n"
        "php artisan key:generate\n"
        "php artisan storage:link\n"
        "php artisan serve"
    )
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(9)
    r_code.font.color.rgb = RGBColor(56, 189, 248)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5. HASIL PENGUJIAN OTOMATIS
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. Hasil Pengujian Otomatis (Automated Testing)")
    r_h5.font.name = "Arial"
    r_h5.font.color.rgb = RGBColor(30, 58, 138)

    p_test = doc.add_paragraph()
    p_test.add_run(
        "Seluruh suite pengujian fitur (Unit & Feature Tests) dieksekusi dengan hasil 100% Lulus (Green):\n"
        "• Total Pengujian: 41 Test Cases\n"
        "• Total Assertion: 126 Assertions\n"
        "• Status: PASS (0 Error, 0 Warning, 0 Failure)\n"
        "• Waktu Eksekusi: ~2.88 Detik"
    )

    # Footer note
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p_ft = doc.add_paragraph()
    p_ft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ft = p_ft.add_run("Dokumen diterbitkan secara otomatis oleh Sistem SIMANTRA BPS\nStatus: Terverifikasi & Terintegrasi ke GitHub Main Branch")
    r_ft.font.name = "Arial"
    r_ft.font.size = Pt(8)
    r_ft.font.italic = True
    r_ft.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(output_docx_path)
    print(f"DOCX created: {output_docx_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)
        doc = word.Documents.Open(abs_docx)
        doc.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
        doc.Close()
        print(f"PDF converted: {abs_pdf}")
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

if __name__ == "__main__":
    docx_file = "Laporan_Pembaruan_SIMANTRA_28_29_Agustus_2026.docx"
    pdf_file = "Laporan_Pembaruan_SIMANTRA_28_29_Agustus_2026.pdf"
    
    create_report_docx(docx_file)
    convert_docx_to_pdf(docx_file, pdf_file)

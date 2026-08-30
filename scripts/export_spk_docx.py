import os
import sys
import json
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_exact_spk_docx(input_docx_path, output_docx_path, data_json_path):
    with open(data_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = docx.Document(input_docx_path)

    nomor_spk = data.get('nomor_spk', '')
    nama_mitra = data.get('nama_mitra', '')
    pekerjaan_mitra = data.get('pekerjaan_mitra', '')
    alamat_mitra = data.get('alamat_mitra', '')
    periode_label = data.get('periode_label', '')
    total_honor = data.get('total_honor', '')
    terbilang_honor = data.get('terbilang_honor', '')
    tahun = str(data.get('tahun', ''))
    hari = data.get('hari', 'Senin')
    tanggal_teks = data.get('tanggal_teks', 'tanggal satu')
    bulan_teks = data.get('bulan_teks', 'bulan Januari')
    tahun_teks = data.get('tahun_teks', 'tahun dua ribu dua puluh lima')
    items = data.get('items', [])

    # 1. Enforce Page Break on Pasal 10
    for p in doc.paragraphs:
        if p.text.strip().startswith('Pasal 10'):
            p.paragraph_format.page_break_before = True

    # 2. Perfect P1 & P2 Alignment with Full Printable Body Width & Left-Flush Indent
    p1 = None
    p2 = None
    for p in doc.paragraphs:
        if '1. Dindin Muldiana' in p.text:
            p1 = p
        if '2. LINA KARLINA' in p.text or ('2. ' in p.text and 'PIHAK KEDUA' in p.text):
            p2 = p

    if p1 is not None and p2 is not None:
        table = doc.add_table(rows=2, cols=3)
        table.autofit = False

        tblPr = table._element.tblPr
        
        # Flush table to the left margin
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '0')
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ['top', 'left', 'bottom', 'right']:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), '0')
            m.set(qn('w:type'), 'dxa')
            tblCellMar.append(m)
        tblPr.append(tblCellMar)

        # Full body width: 17.36 cm
        widths = [docx.shared.Cm(6.8), docx.shared.Cm(0.35), docx.shared.Cm(10.21)]

        # Row 0: Pihak 1
        r0 = table.rows[0]
        for idx, w in enumerate(widths):
            r0.cells[idx].width = w

        p0_0 = r0.cells[0].paragraphs[0]
        p0_0.paragraph_format.space_before = docx.shared.Pt(0)
        p0_0.paragraph_format.space_after = docx.shared.Pt(8)
        p0_0.paragraph_format.line_spacing = 1.15
        run0_0 = p0_0.add_run('1. Dindin Muldiana, S.ST. MP.')
        run0_0.font.name = 'Bookman Old Style'
        run0_0.font.size = docx.shared.Pt(12)

        p0_1 = r0.cells[1].paragraphs[0]
        p0_1.paragraph_format.space_before = docx.shared.Pt(0)
        p0_1.paragraph_format.space_after = docx.shared.Pt(8)
        p0_1.paragraph_format.line_spacing = 1.15
        run0_1 = p0_1.add_run(':')
        run0_1.font.name = 'Bookman Old Style'
        run0_1.font.size = docx.shared.Pt(12)

        p0_2 = r0.cells[2].paragraphs[0]
        p0_2.paragraph_format.space_before = docx.shared.Pt(0)
        p0_2.paragraph_format.space_after = docx.shared.Pt(8)
        p0_2.paragraph_format.line_spacing = 1.15
        p0_2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        run0_2 = p0_2.add_run('Pejabat Pembuat Komitmen Badan Pusat Statistik Kabupaten Tasikmalaya; berkedudukan di  Jl. Raya Timur Singaparna KM4 Cintaraja Tasikmalaya, bertindak untuk dan atas nama Badan Pusat Statistik Kabupaten Tasikmalaya, selanjutnya disebut sebagai ')
        run0_2.font.name = 'Bookman Old Style'
        run0_2.font.size = docx.shared.Pt(12)
        run0_2_b = p0_2.add_run('PIHAK PERTAMA')
        run0_2_b.font.name = 'Bookman Old Style'
        run0_2_b.font.size = docx.shared.Pt(12)
        run0_2_b.bold = True
        run0_2_d = p0_2.add_run('. ')
        run0_2_d.font.name = 'Bookman Old Style'
        run0_2_d.font.size = docx.shared.Pt(12)

        # Row 1: Pihak 2
        r1 = table.rows[1]
        for idx, w in enumerate(widths):
            r1.cells[idx].width = w

        p1_0 = r1.cells[0].paragraphs[0]
        p1_0.paragraph_format.space_before = docx.shared.Pt(0)
        p1_0.paragraph_format.space_after = docx.shared.Pt(0)
        p1_0.paragraph_format.line_spacing = 1.15
        run1_0 = p1_0.add_run(f'2. {nama_mitra}')
        run1_0.font.name = 'Bookman Old Style'
        run1_0.font.size = docx.shared.Pt(12)

        p1_1 = r1.cells[1].paragraphs[0]
        p1_1.paragraph_format.space_before = docx.shared.Pt(0)
        p1_1.paragraph_format.space_after = docx.shared.Pt(0)
        p1_1.paragraph_format.line_spacing = 1.15
        run1_1 = p1_1.add_run(':')
        run1_1.font.name = 'Bookman Old Style'
        run1_1.font.size = docx.shared.Pt(12)

        p1_2 = r1.cells[2].paragraphs[0]
        p1_2.paragraph_format.space_before = docx.shared.Pt(0)
        p1_2.paragraph_format.space_after = docx.shared.Pt(0)
        p1_2.paragraph_format.line_spacing = 1.15
        p1_2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        run1_2 = p1_2.add_run(f'{pekerjaan_mitra}, berkedudukan di {alamat_mitra}, bertindak untuk dan atas nama diri sendiri, selanjutnya disebut ')
        run1_2.font.name = 'Bookman Old Style'
        run1_2.font.size = docx.shared.Pt(12)
        run1_2_b = p1_2.add_run('PIHAK KEDUA')
        run1_2_b.font.name = 'Bookman Old Style'
        run1_2_b.font.size = docx.shared.Pt(12)
        run1_2_b.bold = True
        run1_2_d = p1_2.add_run('.')
        run1_2_d.font.name = 'Bookman Old Style'
        run1_2_d.font.size = docx.shared.Pt(12)

        # Insert table where P1 was
        p1._element.addprevious(table._element)

        # Delete old paragraphs P1, P2 and the middle spacer (preserving top and bottom template spacers)
        p8_elem = p1._element.getnext()
        if p8_elem is not None and p8_elem.tag.endswith('}p'):
            p8_elem.getparent().remove(p8_elem)
        p1._element.getparent().remove(p1._element)
        p2._element.getparent().remove(p2._element)

    # 3. Universal XML Element Replacements
    replacements = [
        # Mail Merge Tanggal Pembukaan
        ('Selasa', hari),
        ('tanggal dua puluh delapan, bulan Februari', f'{tanggal_teks}, {bulan_teks}'),
        ('tahun dua ribu dua puluh empat', tahun_teks),
        ('TAHUN 2024', f'TAHUN {tahun}'),
        ('tahun 2024', f'tahun {tahun}'),

        # Nomor Dokumen
        ('1001/PPK/SPK/03/2024', nomor_spk),

        # Mail Merge Nama di Lampiran & Textbox Shape
        ('MERGEFIELD Nama_Petugas LINA KARLINA', nama_mitra),
        ('MERGEFIELD  Nama_Petugas  LINA KARLINA', nama_mitra),
        ('MERGEFIELD Nama_Petugas', nama_mitra),
        ('MERGEFIELD', ''),
        ('«Nama_Petugas»', nama_mitra),
        ('${Nama_Petugas}', nama_mitra),

        # Nama Mitra di Klausul & Tanda Tangan
        ('LINA KARLINA', nama_mitra),
        ('Nur Azizah Muyassaroh, S.ST', nama_mitra),
        ('${NAMA_MITRA}', nama_mitra),
        ('${NAMA}', nama_mitra),
        ('Lainnya/ Belum Bekerja', pekerjaan_mitra),
        ('${PEKERJAAN}', pekerjaan_mitra),
        ('Kp. Pameungpeuk RT/RW : 24/03 Desa Sukarasa Kec. Salawu', alamat_mitra),
        ('${ALAMAT}', alamat_mitra),

        # Periode & Honor
        ('1 Maret 2024 sampai dengan tanggal 31 Maret 2024', periode_label),
        ('Januari s.d Desember 2024', periode_label),
        ('Rp. 1.160.000 (Satu Juta Seratus Enam Puluh Ribu Rupiah)', f'{total_honor} ({terbilang_honor})'),
        ('Rp. 1.160.000,00 (Satu Juta Seratus Enam Puluh Ribu Rupiah)', f'{total_honor},00 ({terbilang_honor})'),
        ('Satu Juta Seratus Enam Puluh Ribu Rupiah', terbilang_honor),
        ('${TERBILANG}', terbilang_honor),
        ('Rp. 1.160.000, 00', f'{total_honor}, 00'),
        ('Rp. 1.160.000,00', f'{total_honor},00'),
        ('Rp. 1.160.000', total_honor),
        ('${TOTAL_HONOR}', total_honor),
    ]

    # Clean up all red color nodes and apply text replacements
    for elem in list(doc._element.iter()):
        if elem.tag.endswith('}color'):
            for val in elem.attrib.values():
                if 'ff0000' in str(val).lower():
                    if elem.getparent() is not None:
                        elem.getparent().remove(elem)

        if elem.tag.endswith('}t') and elem.text:
            for old_text, new_text in replacements:
                if old_text in elem.text:
                    elem.text = elem.text.replace(old_text, new_text)

    # 4. ANNEX TABLE POPULATION (TARGET LAST TABLE IN DOCUMENT)
    if len(doc.tables) >= 2:
        t1 = doc.tables[-1] # The annex table is always the last table

        def format_cell(cell, text, font_size_pt=9.0, bold=False, italic=False, align=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT):
            cell.text = text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = docx.shared.Pt(0)
            p.paragraph_format.space_after = docx.shared.Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = align
            for r in p.runs:
                r.font.name = 'Bookman Old Style'
                r.font.size = docx.shared.Pt(font_size_pt)
                r.bold = bold
                r.italic = italic

        for i in range(8):
            row_idx = 3 + i
            if row_idx < len(t1.rows):
                row = t1.rows[row_idx]
                if i < len(items):
                    item = items[i]
                    if len(row.cells) >= 8:
                        format_cell(row.cells[0], str(i + 1), font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell(row.cells[1], item.get('nama', ''), font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                        vol_int = int(round(float(item.get('volume', 1) or 1)))
                        format_cell(row.cells[3], str(vol_int), font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell(row.cells[4], item.get('satuan', 'dokumen'), font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell(row.cells[5], item.get('harga_satuan', ''), font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT)
                        format_cell(row.cells[6], item.get('nilai_perjanjian', ''), font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT)
                        format_cell(row.cells[7], item.get('mak', ''), font_size_pt=7.5, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    if len(row.cells) >= 8:
                        format_cell(row.cells[0], str(i + 1), font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell(row.cells[1], '', font_size_pt=9.0)
                        format_cell(row.cells[2], '', font_size_pt=9.0)
                        format_cell(row.cells[3], '', font_size_pt=9.0)
                        format_cell(row.cells[4], '', font_size_pt=9.0)
                        format_cell(row.cells[5], 'Rp. 00', font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT)
                        format_cell(row.cells[6], 'Rp. 00', font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT)
                        format_cell(row.cells[7], '', font_size_pt=7.5)

        if len(t1.rows) > 11:
            total_row = t1.rows[11]
            if len(total_row.cells) >= 7:
                format_cell(total_row.cells[0], f'Terbilang: {terbilang_honor}', font_size_pt=9.0, italic=True, align=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                format_cell(total_row.cells[6], f'{total_honor}, 00', font_size_pt=9.0, align=docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(output_docx_path)

    # Optional direct PDF conversion if requested via --pdf flag
    if len(sys.argv) >= 5 and sys.argv[4] == '--pdf':
        pdf_out_path = output_docx_path.rsplit('.', 1)[0] + '.pdf'
        try:
            import win32com.client, pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            abs_docx = os.path.abspath(output_docx_path)
            abs_pdf = os.path.abspath(pdf_out_path)
            doc_app = word.Documents.Open(abs_docx)
            doc_app.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
            doc_app.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            print(f"SUCCESS_PDF:{abs_pdf}")
        except Exception as e:
            print(f"PDF_CONVERT_ERROR:{str(e)}")

if __name__ == '__main__':
    if len(sys.argv) >= 4:
        generate_exact_spk_docx(sys.argv[1], sys.argv[2], sys.argv[3])
